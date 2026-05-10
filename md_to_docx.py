"""Converte relatorio.md -> relatorio.docx preservando estrutura.

Suporta o subconjunto de markdown usado no relatorio:
- # / ## / ### (titulo, secao, subsecao)
- **negrito** e `codigo` inline
- Blocos de codigo ``` ... ```
- Tabelas com pipes
- Listas com - e 1.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path(__file__).parent / "relatorio.md"
DST = Path(__file__).parent / "relatorio.docx"

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")


def add_inline(paragraph, text: str, *, base_bold: bool = False):
    """Insere texto com formatacao inline (negrito e monoespacado)."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
            if base_bold:
                run.bold = True


def shade_cell(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    rPr.append(shd)


def parse_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, cell_text, base_bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(cell, "D9E2F3")
    doc.add_paragraph()


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for level, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        st = doc.styles[level]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)


def convert():
    md = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(md):
        line = md[i]

        if in_code:
            if line.strip().startswith("```"):
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            in_code = True
            i += 1
            continue

        # Cabecalhos
        if line.startswith("# "):
            p = doc.add_heading(level=0)
            add_inline(p, line[2:].strip())
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_heading(level=1)
            add_inline(p, line[3:].strip())
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_heading(level=2)
            add_inline(p, line[4:].strip())
            i += 1
            continue

        # Tabela: linha com | seguida por linha separadora
        if "|" in line and i + 1 < len(md) and TABLE_SEP_RE.match(md[i + 1]):
            header = parse_table_row(line)
            rows = [header]
            i += 2
            while i < len(md) and "|" in md[i] and md[i].strip():
                rows.append(parse_table_row(md[i]))
                i += 1
            add_table(doc, rows)
            continue

        # Lista nao-ordenada
        if re.match(r"^\s*-\s+", line):
            indent_spaces = len(line) - len(line.lstrip())
            content = re.sub(r"^\s*-\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            if indent_spaces >= 2:
                p.paragraph_format.left_indent = Cm(0.6 + 0.6 * (indent_spaces // 2))
            add_inline(p, content)
            i += 1
            continue

        # Lista ordenada
        if re.match(r"^\s*\d+\.\s+", line):
            content = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, content)
            i += 1
            continue

        # Linha em branco -> espaco
        if not line.strip():
            i += 1
            continue

        # Paragrafo comum (acumula linhas consecutivas)
        buf = [line]
        i += 1
        while (
            i < len(md)
            and md[i].strip()
            and not md[i].startswith("#")
            and not md[i].strip().startswith("```")
            and not re.match(r"^\s*-\s+", md[i])
            and not re.match(r"^\s*\d+\.\s+", md[i])
            and not ("|" in md[i] and i + 1 < len(md) and TABLE_SEP_RE.match(md[i + 1]))
        ):
            buf.append(md[i])
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(s.strip() for s in buf))

    doc.save(DST)
    print(f"OK: {DST}")


if __name__ == "__main__":
    convert()
